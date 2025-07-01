import streamlit as st
import sys
import os
import pdfplumber
from openai import OpenAI
from datetime import datetime
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Pinecone as PineconeVectorStore
from langchain.chains import RetrievalQA
from langchain.chat_models import ChatOpenAI
from io import BytesIO
from docx import Document
from docx.shared import Pt
from pinecone import Pinecone


# 상위 디렉토리의 utils 모듈 import를 위한 경로 추가
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# 페이지 설정
st.set_page_config(
    page_title="AI_리포트 생성기",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 커스텀 CSS
st.markdown("""
<style>
    .info-container {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 25px;
        border-radius: 15px;
        color: white;
        margin: 20px 0;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    .data-source-card {
        background: white;
        padding: 20px;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        margin: 15px 0;
        border-left: 5px solid #1f77b4;
    }
    .update-card {
        background: linear-gradient(135deg, #74b9ff 0%, #0984e3 100%);
        padding: 15px;
        border-radius: 10px;
        color: white;
        margin: 10px 0;
    }
    .system-info {
        background: linear-gradient(135deg, #00b894 0%, #00a085 100%);
        padding: 20px;
        border-radius: 15px;
        color: white;
        margin: 15px 0;
    }
    .guide-section {
        background: #f8f9fa;
        padding: 20px;
        border-radius: 10px;
        margin: 15px 0;
        border: 1px solid #e9ecef;
    }
</style>
""", unsafe_allow_html=True)

# 타이틀
st.markdown('<h1 class="main-header">📄 AI 기반 보고서 생성기</h1>', unsafe_allow_html=True)

# OpenAI API Key 로딩
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")

pinecone_api_key = os.getenv("PINECONE_API_KEY")
pinecone_env=os.getenv("PINECONE_ENV")
pc = Pinecone(api_key=pinecone_api_key)
index_name = "carbone-index"

# PDF 텍스트 추출
@st.cache_data
def extract_text_from_pdf(uploaded_file):
    with pdfplumber.open(uploaded_file) as pdf:
        text = ""
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

# 목차 추출 함수
def extract_table_of_contents(text):
    prompt = f"""
다음 문서에서 **목차(차례)**에 해당하는 부분만 정확히 추출해 주세요.
- 숫자나 로마자, 제목 패턴을 이용해 목차 항목만 뽑아주세요.
- 본문 내용은 포함하지 말고, 목차 구조만 출력하세요.

문서 내용:
{text[:4000]}
"""
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "당신은 문서 구조에서 목차만 정확히 추출하는 AI입니다."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3
    )
    return response.choices[0].message.content.strip()

# 문서 전체 구조 요약
def summarize_template_structure(text):
    prompt = f"""
다음 문서의 형식(보고서 구조, 제목 스타일, 구성 흐름 등)을 간단히 요약해 주세요.
- 문서가 어떤 형식으로 작성되어 있는지 설명해주세요.
- 목차, 본문 구성, 언어 톤 등을 포함해 형식을 분석해주세요.

문서 내용:
{text[:4000]}
"""
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "당신은 문서 형식을 분석하고 요약하는 AI입니다."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.5
    )
    return response.choices[0].message.content.strip()

#Pinecone 기반 RAG용 DB 생성
@st.cache_resource
def create_vector_store(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    docs = text_splitter.create_documents([text])
    embeddings = OpenAIEmbeddings()

    # ✅ 인덱스 존재 여부 확인 → 없으면 생성
    # ✅ 인덱스가 없으면 생성
    if index_name not in pc.list_indexes().names():
        from pinecone import ServerlessSpec
        pc.create_index(
            name=index_name,
            dimension=1536,
            metric="cosine",
            spec=ServerlessSpec(cloud="aws", region="us-east-1")
        )

    vector_store = PineconeVectorStore.from_documents(docs, embeddings, index_name=index_name)
    return vector_store


#Word 보고서 생성 함수
def generate_docx_report(text, topic):
    doc = Document()
    doc.add_heading(topic, level=1)

    for paragraph in text.split("\n"):
        if paragraph.strip():
            p = doc.add_paragraph(paragraph.strip())
            p.style.font.size=Pt(11)
            p.paragraph_format.line_spacing = 1.5

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 보고서 생성
def generate_report_with_rag(topic, vector_store):
    retriever = vector_store.as_retriever(search_kwargs={"k": 5})
    llm = ChatOpenAI(temperature=0.7, model_name="gpt-4")
    qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever)

    prompt = f"""
'{topic}'에 대한 보고서를 작성해 주세요. 
- 참고 문서를 기반으로 구성과 형식을 따라주세요.
- 목차 구성, 말투, 분석 방향을 유지하고, 새로운 주제에 맞게 내용을 작성하세요.
- 참고 문서는 참고만 할뿐이지, 거기의 내용을 가지고와서는 안됩니다.
"""
    return qa.run(prompt)

# # 주제 기반 보고서 생성
# def generate_report_based_on_template(topic, template_text):
#     prompt = f"""
# 아래 문서 형식을 참고하여, 새로운 주제 '{topic}'에 대해 동일한 형식의 보고서를 작성해 주세요.
#
# - 문서 형식(목차, 구성, 말투 등)은 그대로 유지하되,
# - 내용은 '{topic}'을 기반으로 완전히 새롭게 작성해 주세요.
#
# 📄 참고 문서:
# {template_text[:4000]}
# """
#     response = client.chat.completions.create(
#         model="gpt-4",
#         messages=[
#             {"role": "system", "content": "당신은 보고서 형식을 학습해 새로운 주제에 맞춰 작성하는 AI입니다."},
#             {"role": "user", "content": prompt}
#         ],
#         temperature=0.7
#     )
#     return response.choices[0].message.content.strip()

uploaded_file = st.file_uploader("📎 PDF 문서를 업로드하세요", type=["pdf"])

if uploaded_file:
    with st.spinner("문서 분석 중..."):
        extracted_text = extract_text_from_pdf(uploaded_file)
        toc = extract_table_of_contents(extracted_text)
        structure_summary = summarize_template_structure(extracted_text)
        vector_store = create_vector_store(extracted_text)

    st.subheader("🧾 문서 목차 자동 추출")
    st.code(toc, language="markdown")

    st.subheader("📚 문서 형식 요약")
    st.markdown(structure_summary)

    topic = st.text_input("📝 새로 작성할 보고서 주제를 입력하세요", placeholder="예: 탄소중립 추진 전략")

    if topic:
        with st.spinner("📄 새로운 보고서를 작성하는 중입니다..."):
            generated_report = generate_report_with_rag(topic, vector_store)
            docx_file = generate_docx_report(generated_report, topic)

        st.success("✅ 보고서 생성 완료!")
        st.download_button(
            label="📥 Word 형식으로 보고서 다운로드 (.docx)",
            data=docx_file,
            file_name=f"{topic}_보고서.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        st.text_area("📄 생성된 보고서 미리보기", generated_report, height=500)


# topic = st.text_input("📝 새로 작성할 보고서 주제를 입력하세요", placeholder="예: 탄소중립 추진 전략")
#
# # Step 2: 주제 입력 후 보고서 생성
# if uploaded_file and topic:
#     with st.spinner("📄 새로운 보고서를 작성하는 중입니다..."):
#         generated_report = generate_report_based_on_template(topic, extracted_text)
#     st.success("✅ 보고서 생성 완료!")
#
#     st.download_button(
#         "📥 보고서 다운로드",
#         generated_report,
#         file_name=f"{topic}_보고서.txt",
#         mime="text/plain"
#     )
#
#     st.text_area("📄 생성된 보고서 미리보기", generated_report, height=500)